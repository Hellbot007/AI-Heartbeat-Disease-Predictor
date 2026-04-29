from docx import Document

def extract_docx_info(filepath, outpath):
    try:
        doc = Document(filepath)
        with open(outpath, "w", encoding="utf-8") as f:
            f.write("--- PARAGRAPHS ---\n")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    f.write(f"[{i}]: {para.text}\n")
            
            f.write("\n--- TABLES ---\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}:\n")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    f.write(" | ".join(row_data) + "\n")
                f.write("-" * 20 + "\n")
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    extract_docx_info("AI_Heartbeat_Disease_Predictor_Report (1).docx", "tmp/docx_content.txt")
